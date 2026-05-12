"""Chat export endpoint — PDF, Markdown, DOCX."""

from fastapi import APIRouter, Depends, Query
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
import io
from datetime import datetime

from backend.db.session import get_db
from backend.models.entities import Conversation, Message

export_router = APIRouter()


async def _get_messages(session_id: str, db: AsyncSession) -> list[dict]:
    """Retrieve all messages for a session."""
    conv = (await db.execute(
        select(Conversation).where(Conversation.session_id == session_id)
    )).scalar_one_or_none()
    
    if not conv:
        return []
    
    stmt = select(Message).where(Message.conversation_id == conv.id).order_by(Message.created_at)
    result = await db.execute(stmt)
    return [{"role": m.role, "content": m.content, "time": m.created_at.isoformat()} for m in result.scalars().all()]


def _render_markdown(messages: list[dict], session_id: str) -> str:
    """Render chat history as Markdown."""
    lines = [
        f"# TiO Chat Export",
        f"**Session:** `{session_id}`",
        f"**Exported:** {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}",
        f"**Messages:** {len(messages)}",
        "",
        "---",
        "",
    ]
    for msg in messages:
        role = "🧑 User" if msg["role"] == "user" else "🤖 Assistant"
        lines.append(f"### {role}")
        lines.append(msg["content"])
        lines.append("")
    return "\n".join(lines)


@export_router.get("/chat/export/{session_id}")
async def export_chat(
    session_id: str,
    format: str = Query("md", regex="^(pdf|md|docx)$"),
    db: AsyncSession = Depends(get_db)
):
    messages = await _get_messages(session_id, db)
    
    if format == "md":
        content = _render_markdown(messages, session_id)
        return StreamingResponse(
            io.BytesIO(content.encode("utf-8")),
            media_type="text/markdown",
            headers={"Content-Disposition": f"attachment; filename=tio-chat-{session_id[:8]}.md"}
        )
    
    elif format == "pdf":
        try:
            from fpdf import FPDF
        except ImportError:
            # Fallback: return markdown if fpdf2 not installed
            content = _render_markdown(messages, session_id)
            return StreamingResponse(
                io.BytesIO(content.encode("utf-8")),
                media_type="text/markdown",
                headers={"Content-Disposition": f"attachment; filename=tio-chat-{session_id[:8]}.md"}
            )
        
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=20)
        
        # Title
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 12, "TiO Chat Export", ln=True)
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(100, 100, 100)
        pdf.cell(0, 6, f"Session: {session_id[:16]}... | {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}", ln=True)
        pdf.ln(8)
        
        for msg in messages:
            # Role header
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(0, 0, 0)
            role = "User" if msg["role"] == "user" else "Assistant"
            pdf.cell(0, 8, role, ln=True)
            
            # Content
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(40, 40, 40)
            # Handle unicode safely
            safe_content = msg["content"].encode("latin-1", "replace").decode("latin-1")
            pdf.multi_cell(0, 5, safe_content)
            pdf.ln(4)
        
        buf = io.BytesIO(pdf.output())
        return StreamingResponse(
            buf,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=tio-chat-{session_id[:8]}.pdf"}
        )
    
    elif format == "docx":
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
        except ImportError:
            content = _render_markdown(messages, session_id)
            return StreamingResponse(
                io.BytesIO(content.encode("utf-8")),
                media_type="text/markdown",
                headers={"Content-Disposition": f"attachment; filename=tio-chat-{session_id[:8]}.md"}
            )
        
        doc = Document()
        doc.add_heading("TiO Chat Export", level=1)
        doc.add_paragraph(f"Session: {session_id[:16]}... | {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
        
        for msg in messages:
            role = "User" if msg["role"] == "user" else "Assistant"
            p = doc.add_paragraph()
            run = p.add_run(f"{role}: ")
            run.bold = True
            run.font.size = Pt(11)
            p.add_run(msg["content"])
        
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=tio-chat-{session_id[:8]}.docx"}
        )
