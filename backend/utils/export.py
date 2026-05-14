import io
from fpdf import FPDF
from docx import Document
from datetime import datetime

class ExportService:
    @staticmethod
    def to_markdown(messages: list[dict], session_id: str) -> str:
        md = f"# TiO Chat Export\n"
        md += f"**Session ID:** {session_id}\n"
        md += f"**Date:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
        md += "---\n\n"
        
        for msg in messages:
            role = msg.get("role", "unknown").upper()
            content = msg.get("content", "")
            md += f"### {role}\n{content}\n\n"
            if msg.get("sources"):
                md += "**Sources:**\n"
                for s in msg["sources"]:
                    md += f"- {s.get('document', 'Unknown')}\n"
                md += "\n"
        return md

    @staticmethod
    def to_pdf(messages: list[dict], session_id: str) -> bytes:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, "TiO Chat Export", ln=True, align="C")
        pdf.set_font("Arial", "", 10)
        pdf.cell(0, 10, f"Session: {session_id}", ln=True, align="C")
        pdf.cell(0, 10, f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True, align="C")
        pdf.ln(10)

        for msg in messages:
            role = msg.get("role", "unknown").upper()
            content = msg.get("content", "")
            
            pdf.set_font("Arial", "B", 12)
            pdf.set_text_color(0, 198, 255) if role == "ASSISTANT" else pdf.set_text_color(0, 0, 0)
            pdf.cell(0, 10, f"{role}:", ln=True)
            
            pdf.set_font("Arial", "", 11)
            pdf.set_text_color(50, 50, 50)
            # Use multi_cell for wrapping text
            pdf.multi_cell(0, 7, content)
            pdf.ln(5)
            
            if msg.get("sources"):
                pdf.set_font("Arial", "I", 9)
                pdf.cell(0, 5, "Sources:", ln=True)
                for s in msg["sources"]:
                    pdf.cell(0, 5, f"- {s.get('document', 'Unknown')}", ln=True)
                pdf.ln(5)
            
            pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 190, pdf.get_y())
            pdf.ln(5)

        return pdf.output(dest='S')

    @staticmethod
    def to_docx(messages: list[dict], session_id: str) -> bytes:
        doc = Document()
        doc.add_heading('TiO Chat Export', 0)
        doc.add_paragraph(f"Session: {session_id}")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        for msg in messages:
            role = msg.get("role", "unknown").upper()
            content = msg.get("content", "")
            
            p = doc.add_paragraph()
            run = p.add_run(f"{role}:")
            run.bold = True
            
            doc.add_paragraph(content)
            
            if msg.get("sources"):
                doc.add_paragraph("Sources:", style='List Bullet')
                for s in msg["sources"]:
                    doc.add_paragraph(s.get('document', 'Unknown'), style='List Bullet 2')
            
            doc.add_page_break() # Optional: page break per exchange or just separator
            
        file_stream = io.BytesIO()
        doc.save(file_stream)
        return file_stream.getvalue()
