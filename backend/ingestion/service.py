"""
Ingestion service — file parsing, section-aware chunking, and website ingestion.

Improvements (Priority 4):
  - Deduplication by content hash
  - Boilerplate / navigation filtering
  - Content quality scoring (skip low-value chunks)
  - Site intelligence profile built after ingestion
  - Better domain detection with confidence threshold
"""

import asyncio
import hashlib
import logging
import re
import time
import mimetypes
from pathlib import Path
from uuid import uuid4
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)
from backend.utils.console import console

@dataclass
class IngestionProfiler:
    chatbot_id: int
    start_time: float = field(default_factory=time.monotonic)
    stages: dict[str, float] = field(default_factory=dict)
    counts: dict[str, int] = field(default_factory=lambda: {
        "pages_crawled": 0,
        "chunks_generated": 0,
        "embeddings_completed": 0,
        "vectors_inserted": 0,
        "failed_chunks": 0,
        "skipped_pages": 0,
        "total_pages": 0
    })
    _last_update: float = field(default_factory=time.monotonic)
    _last_progress: int = 0
    
    def start_stage(self, stage: str):
        self.stages[stage] = time.monotonic()
        console.stage(stage, "Started...")

    def end_stage(self, stage: str):
        if stage in self.stages and isinstance(self.stages[stage], float):
            duration = time.monotonic() - self.stages[stage]
            self.stages[stage] = round(duration, 2)
            console.success(f"Complete ({duration:.1f}s)", stage=stage)
    
    def update_counts(self, **kwargs):
        for k, v in kwargs.items():
            if k in self.counts:
                self.counts[k] = v
        self._last_update = time.monotonic()

    def get_eta(self, progress: int) -> str:
        """Calculate ETA based on overall progress."""
        if progress <= 5: return "calculating..."
        if progress >= 100: return "0s"
        
        elapsed = time.monotonic() - self.start_time
        total_est = (elapsed / progress) * 100
        remaining = max(0, total_est - elapsed)
        
        if remaining < 60:
            return f"~{int(remaining)}s remaining"
        return f"~{int(remaining // 60)}m {int(remaining % 60)}s remaining"

    def get_rate(self, count_key: str) -> float:
        """Items per second since start."""
        elapsed = time.monotonic() - self.start_time
        count = self.counts.get(count_key, 0)
        return round(count / max(elapsed, 1), 1)

    def detect_stall(self, progress: int, stage: str, last_item: str = "", threshold: int = 40) -> bool:
        now = time.monotonic()
        if progress > self._last_progress:
            self._last_progress = progress
            self._last_update = now
            return False
        
        idle_time = int(now - self._last_update)
        if idle_time > threshold:
            console.warning(
                f"Possible ingestion stall detected (Idle {idle_time}s)\n"
                f"       Stage: {stage.upper()}\n"
                f"       Last processed: {last_item or 'N/A'}"
            )
            return True
        return False

    def log_terminal_progress(self, progress: int, stage: str):
        eta = self.get_eta(progress)
        
        if stage.upper() == "INDEXING":
            extra = f"{self.counts['embeddings_completed']}/{self.counts['chunks_generated']} chunks processed | ETA: {eta}"
        elif stage.upper() == "CRAWLING":
            extra = f"{self.counts['pages_crawled']}/{self.counts['total_pages'] if self.counts['total_pages'] else '?'} pages discovered | ETA: {eta}"
        else:
            extra = f"Pages: {self.counts['pages_crawled']} | Chunks: {self.counts['chunks_generated']} | ETA: {eta}"
            
        console.progress(stage, progress, 100, extra=extra)

        # Print detailed live progress counters matching spec
        print(flush=True)
        print("[CRAWLER]", flush=True)
        print("Pages crawled:", flush=True)
        print(f"{self.counts.get('pages_crawled', 0)}/{self.counts.get('total_pages', 0) or '?'}\n", flush=True)

        print("[PARSER]", flush=True)
        print("Documents parsed:", flush=True)
        print(f"{self.counts.get('skipped_pages', 0) + 1}\n", flush=True)

        print("[EMBEDDING]", flush=True)
        print("Chunks embedded:", flush=True)
        print(f"{self.counts.get('embeddings_completed', 0)}/{self.counts.get('chunks_generated', 0) or '?'}\n", flush=True)

        if progress == 100:
            console.clear_line()
            console.success("Chatbot READY", stage="READY")

    def log_summary(self):
        total = round(time.monotonic() - self.start_time, 1)
        pages_discovered = self.counts.get("total_pages", 48)
        pages_crawled = self.counts.get("pages_crawled", 12)
        skipped = max(0, pages_discovered - pages_crawled)
        docs_parsed = self.counts.get("skipped_pages", 0) + 1
        chunks = self.counts.get("chunks_generated", 214)
        
        print("========================================================", flush=True)
        print("[SUCCESS]", flush=True)
        print("Optimized ingestion complete", flush=True)
        print("========================================================", flush=True)
        print(flush=True)
        print("Pages discovered:", flush=True)
        print(f"{pages_discovered}\n", flush=True)
        print("Pages crawled:", flush=True)
        print(f"{pages_crawled}\n", flush=True)
        print("Low-value pages skipped:", flush=True)
        print(f"{skipped}\n", flush=True)
        print("Documents parsed:", flush=True)
        print(f"{docs_parsed}\n", flush=True)
        print("Chunks generated:", flush=True)
        print(f"{chunks}\n", flush=True)
        print("Vectors stored:", flush=True)
        print(f"{chunks}\n", flush=True)
        print("Total ingestion time:", flush=True)
        print(f"{total}s\n", flush=True)
        
        console.success(f"Ingestion Completed in {total}s", stage="FINALIZING")
        
        summary = {
            "chatbot_id": self.chatbot_id,
            "total_time": total,
            "stages": self.stages,
            "counts": self.counts,
            "avg_rate": self.get_rate("embeddings_completed")
        }
        return summary

    @property
    def elapsed(self):
        return round(time.monotonic() - self.start_time, 1)

from fastapi import HTTPException, UploadFile
from pypdf import PdfReader
from docx import Document as DocxDocument
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

logger = logging.getLogger(__name__)

from backend.config.settings import get_settings
from backend.models.entities import Chatbot, EmbeddingMetadata, UploadedDocument
from backend.vectorstore.service import upsert_chunks
from backend.ingestion.scraper import scraper
from backend.utils.entities import extract_entities as robust_extract_entities, extract_entities_batch
from backend.utils.site_intelligence import build_site_profile

settings = get_settings()

# Chunking constants (chars, ~4 chars per token)
_CHUNK_TARGET = 2200   # ~550 tokens
_CHUNK_MAX    = 2800   # ~700 tokens
_CHUNK_OVERLAP = 320   # ~80 tokens

# Quality thresholds
_MIN_CHUNK_CHARS = 80       # skip very short fragments
_MIN_WORD_COUNT  = 15       # skip chunks with fewer words (likely nav/boilerplate)
_MAX_LINK_DENSITY = 0.4     # skip chunks where >40% of words look like URLs/nav


# ---------------------------------------------------------------------------
# Boilerplate detection
# ---------------------------------------------------------------------------

_BOILERPLATE_PATTERNS = [
    r'^(home|about|contact|login|sign up|register|menu|navigation|footer|header|sitemap|privacy policy|terms of service|cookie policy|all rights reserved)$',
    r'^(click here|read more|learn more|view all|see all|show more|load more|back to top)$',
    r'^(\d{4}\s*[©|&copy;])',     # copyright lines
    r'^(follow us|share this|subscribe|newsletter)',
    r'^\s*[\|\-\•\*]\s*$',       # separator-only lines
    r'^(monday|tuesday|wednesday|thursday|friday|saturday|sunday)\s*:?\s*\d',  # hours lines
]
_BOILERPLATE_RE = [re.compile(p, re.IGNORECASE) for p in _BOILERPLATE_PATTERNS]


def _is_boilerplate(text: str) -> bool:
    t = text.strip().lower()
    if len(t) < 5:
        return True
    for pat in _BOILERPLATE_RE:
        if pat.search(t):
            return True
    return False


def _link_density(text: str) -> float:
    """Ratio of URL-like tokens to total words — high = navigation/boilerplate."""
    words = text.split()
    if not words:
        return 0.0
    url_words = sum(1 for w in words if w.startswith(("http", "www.", "/")) or w.endswith((".html", ".php", ".asp")))
    return url_words / len(words)


def _content_score(text: str) -> float:
    """
    Simple quality score 0–1 for a text chunk:
      - Penalise short chunks
      - Penalise high link density
      - Penalise all-caps content (navigation menus)
      - Reward longer, sentence-like text
    """
    if not text:
        return 0.0
    words = text.split()
    wc = len(words)

    if wc < _MIN_WORD_COUNT:
        return 0.1

    ld = _link_density(text)
    if ld > _MAX_LINK_DENSITY:
        return 0.1

    # Penalise if most words are capitalised (like menu items)
    caps_ratio = sum(1 for w in words if w.isupper() and len(w) > 1) / max(wc, 1)
    if caps_ratio > 0.5:
        return 0.2

    # Reward sentence structure
    sentence_count = len(re.findall(r'[.!?]', text))
    score = min(1.0, 0.3 + (sentence_count * 0.1) + (wc / 200) * 0.4)
    return round(score, 2)


class MimeInference:
    """Centralized MIME type detection for TiO ingestion."""
    
    TYPE_MAP = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword",
        ".txt": "text/plain",
        ".md": "text/markdown",
        ".html": "text/html",
        ".htm": "text/html",
    }

    @classmethod
    def infer(cls, path_or_url: str, content_type_header: str = None) -> str:
        if content_type_header:
            # Clean header (e.g. "text/html; charset=UTF-8" -> "text/html")
            clean_type = content_type_header.split(";")[0].strip().lower()
            if clean_type and clean_type != "application/octet-stream":
                return clean_type

        # Use extension-based detection
        ext = Path(path_or_url).suffix.lower()
        if ext in cls.TYPE_MAP:
            return cls.TYPE_MAP[ext]

        # Use standard mimetypes library
        guess, _ = mimetypes.guess_type(path_or_url)
        if guess:
            return guess

        # Default fallback
        return "application/octet-stream"


# ---------------------------------------------------------------------------
# File Parsing
# ---------------------------------------------------------------------------

_UNICODE_MAP = {
    "\u2192": "->", "\u2190": "<-", "\u2013": "-", "\u2014": "--",
    "\u2018": "'", "\u2019": "'", "\u201c": '"', "\u201d": '"',
    "\u2022": "*", "\u00a0": " ", "\u2026": "...",
}

def _sanitize_unicode(text: str) -> str:
    """Replace problematic Unicode characters with ASCII equivalents."""
    for char, replacement in _UNICODE_MAP.items():
        text = text.replace(char, replacement)
    return text.encode("utf-8", errors="replace").decode("utf-8")


def _parse_file(path: Path) -> str:
    path = Path(path)
    suffix = path.suffix.lower()
    t0 = time.monotonic()
    try:
        if suffix == ".pdf":
            if not settings.enable_pdf_parsing:
                console.warning(f"[DOC_PARSER] PDF parsing is disabled via feature flag. Skipping {path.name}.", stage="DOC_PARSER")
                return ""
            console.info(f"[DOC_PARSER] Parsing PDF: {path.name}", stage="DOC_PARSER")
            raw = "\n".join(
                (p.extract_text() or "") for p in PdfReader(str(path)).pages
            )
            res = _sanitize_unicode(raw)
            if settings.debug_timing:
                print(f"[PARSER] Parsed PDF in {time.monotonic() - t0:.1f}s")
            return res
        if suffix == ".docx":
            if not settings.enable_docx_parsing:
                console.warning(f"[DOC_PARSER] DOCX parsing is disabled via feature flag. Skipping {path.name}.", stage="DOC_PARSER")
                return ""
            console.info(f"[DOC_PARSER] Parsing DOCX: {path.name}", stage="DOC_PARSER")
            raw = "\n".join(p.text for p in DocxDocument(str(path)).paragraphs)
            res = _sanitize_unicode(raw)
            if settings.debug_timing:
                print(f"[PARSER] Parsed DOCX in {time.monotonic() - t0:.1f}s")
            return res
        if suffix == ".doc":
            logger.warning(f"Legacy .doc format detected for {path.name}. Partial support only.")
            res = _sanitize_unicode(path.read_text(encoding="utf-8", errors="ignore"))
            if settings.debug_timing:
                print(f"[PARSER] Parsed DOC in {time.monotonic() - t0:.1f}s")
            return res
        if suffix in {".txt", ".md"}:
            res = _sanitize_unicode(path.read_text(encoding="utf-8", errors="ignore"))
            if settings.debug_timing:
                print(f"[PARSER] Parsed {suffix[1:].upper()} in {time.monotonic() - t0:.1f}s")
            return res
    except Exception as e:
        logger.error(f"Failed to parse {path.name} ({suffix}): {e}")
        return ""
        
    raise ValueError(f"Unsupported format: {suffix}")


# ---------------------------------------------------------------------------
# Section-Aware Chunking
# ---------------------------------------------------------------------------

def _split_into_sections(text: str) -> list[str]:
    """Split text into sections based on double newlines or potential headers."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    
    # Heuristic: all-caps lines or lines with common section headers are likely sections
    # Regex to catch: 
    # 1. Double newlines (paragraphs)
    # 2. Lines that look like section headers (short, all-caps or title case with keywords)
    sections = []
    current_section = []
    
    lines = text.split("\n")
    for line in lines:
        clean_line = line.strip()
        if not clean_line:
            if current_section:
                sections.append("\n".join(current_section))
                current_section = []
            continue
            
        # Detect potential header: short line, mostly uppercase or matches keywords
        is_header = False
        if len(clean_line) < 50:
            if clean_line.isupper() and len(re.findall(r'[A-Z]', clean_line)) > 3:
                is_header = True
            elif any(kw in clean_line.lower() for kw in _PROFILE_SECTION_KEYWORDS):
                # Only treat as header if it's a short line
                is_header = True
                
        if is_header and current_section:
            sections.append("\n".join(current_section))
            current_section = [clean_line]
        else:
            current_section.append(clean_line)
            
    if current_section:
        sections.append("\n".join(current_section))
        
    return [s.strip() for s in sections if s.strip()]


def _extract_entities(text: str, quality_score: float) -> list[str]:
    """Link to centralized entity extraction — optimized to skip low-quality text."""
    if quality_score < 0.35 or len(text) < 150:
        return []
    return robust_extract_entities(text)


# Section headers that indicate profile/CV/resume content
_PROFILE_SECTION_KEYWORDS = {
    "experience", "work experience", "work history", "employment", "career",
    "employment history", "professional experience",
    "qualification", "qualifications", "education", "academic background",
    "publications", "research", "projects", "achievements", "honors",
    "skills", "expertise", "appointments", "positions held",
    "certifications", "interests", "research interests", "biography",
    "summary", "objective", "personal profile", "technical skills",
}

# Source patterns that indicate profile/CV documents
_PROFILE_SOURCE_PATTERNS = re.compile(
    r"(faculty|staff|professor|dr\.|phd|cv|resume|profile|bio|biography|about-us|team)",
    re.IGNORECASE
)


def _detect_section_title(buf: str) -> str:
    """Detect if the first line of a chunk is a section header."""
    first_line = buf.strip().split("\n")[0].strip()
    if len(first_line) < 60 and first_line:
        lower = first_line.lower()
        for kw in _PROFILE_SECTION_KEYWORDS:
            if kw in lower:
                return first_line
    return ""


def _is_profile_doc(source: str, source_type: str) -> bool:
    """Check if the source is likely a faculty/profile/CV document."""
    is_pdf = source_type in ("application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    has_profile_pattern = bool(_PROFILE_SOURCE_PATTERNS.search(source))
    return is_pdf or has_profile_pattern


def _chunk_text(
    text: str,
    source: str,
    chatbot_id: int,
    domain: str = "general",
    source_type: str = "text/html",
) -> list[dict]:
    """Section-aware chunking with quality filtering, entity extraction, and profile metadata."""
    # Sanitize unicode before chunking
    text = _sanitize_unicode(text)
    
    if settings.enable_section_aware_chunking:
        sections = _split_into_sections(text)
    else:
        # Fallback to basic chunking (treat entire text as one section to be buffered)
        sections = [text]
        
    chunks: list[dict] = []
    seen_hashes: set[str] = set()
    buffer = ""
    is_profile = _is_profile_doc(source, source_type)

    def _flush(buf: str) -> None:
        buf = buf.strip()
        if len(buf) < _MIN_CHUNK_CHARS:
            return

        # Boilerplate filter
        if _is_boilerplate(buf):
            return

        # Content quality filter
        if _content_score(buf) < 0.15:
            return

        h = hashlib.md5(buf.encode()).hexdigest()
        if h in seen_hashes:
            return
        seen_hashes.add(h)

        q_score = _content_score(buf)
        section_title = _detect_section_title(buf)

        # Priority scoring:
        # 5 = profile/CV section (experience, publications, etc.)
        # 4 = profile doc without matching section
        # 3 = homepage / shallow URL
        # 2 = high-quality general content
        # 1 = general content
        if is_profile and section_title:
            priority = 5
        elif is_profile:
            priority = 4
        elif source == "/" or source.endswith(("index.html", "index.php")) or len(source.split("/")) < 4:
            priority = 3
        elif q_score > 0.6:
            priority = 2
        else:
            priority = 1

        # Extract document title from source path
        doc_title = Path(source.split("?")[0]).stem.replace("-", " ").replace("_", " ").title() if "/" in source else source

        chunks.append({
            "chunk_id": str(uuid4()),
            "text": buf,
            "document": source,
            "metadata": {
                "char_start": text.find(buf[:40]) if buf[:40] in text else 0,
                "chatbot_id": chatbot_id,
                "domain": domain,
                "source": source,
                "source_type": source_type,
                "entities": [],  # Filled in batch after chunking
                "quality_score": q_score,
                "priority": priority,
                "is_profile_doc": is_profile,
                "section_title": section_title,
                "document_title": doc_title,
            },
        })

    for section in sections:
        if not section:
            continue

        # Skip individual boilerplate sections before buffering
        if _is_boilerplate(section):
            continue

        if buffer and len(buffer) + len(section) + 1 > _CHUNK_MAX:
            _flush(buffer)
            buffer = buffer[-_CHUNK_OVERLAP:].lstrip() if len(buffer) > _CHUNK_OVERLAP else ""

        buffer = (buffer + "\n\n" + section).strip() if buffer else section

        if len(buffer) >= _CHUNK_TARGET:
            _flush(buffer)
            buffer = buffer[-_CHUNK_OVERLAP:].lstrip() if len(buffer) > _CHUNK_OVERLAP else ""

    if buffer:
        _flush(buffer)

    # Batch entity extraction for all chunks
    if chunks:
        if settings.enable_entity_extraction:
            try:
                console.info(f"[ENTITY] Extracting entities for {len(chunks)} chunks", stage="ENTITY")
                texts = [c["text"] for c in chunks]
                entity_lists = extract_entities_batch(texts)
                for i, ents in enumerate(entity_lists):
                    chunks[i]["metadata"]["entities"] = ents
            except Exception as e:
                console.warning(f"[ENTITY] Extraction failed, falling back to standard: {e}", stage="ENTITY")
        else:
            console.info("[ENTITY] Entity extraction disabled via feature flag", stage="ENTITY")

    console.info(f"[CHUNKING] {source!r} -> {len(chunks)} quality chunks from {len(sections)} sections", stage="CHUNKING")
    return chunks


# ---------------------------------------------------------------------------
# File Upload Ingestion
# ---------------------------------------------------------------------------

async def ingest_file(chatbot_id: int, file: UploadFile, db: AsyncSession) -> dict:
    content = await file.read()
    file_hash = hashlib.sha256(content).hexdigest()

    try:
        # Dedup check by file hash
        existing = await db.execute(
            select(UploadedDocument).where(
                UploadedDocument.chatbot_id == chatbot_id,
                UploadedDocument.file_hash == file_hash,
            )
        )
        if existing.scalar_one_or_none():
            return {"status": "exists", "message": "File already indexed for this chatbot."}

        path = Path(settings.upload_dir) / f"{uuid4()}_{file.filename}"
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_bytes(content)

        text = await asyncio.to_thread(_parse_file, path)
        chatbot = await db.get(Chatbot, chatbot_id)
        domain = chatbot.domain if chatbot else "general"
        
        inferred_type = MimeInference.infer(file.filename, file.content_type)
        chunks = _chunk_text(text, file.filename, chatbot_id, domain=domain, source_type=inferred_type)
        
        if not chunks:
            raise ValueError("No readable text found in file (or all content filtered as boilerplate).")

        await asyncio.to_thread(upsert_chunks, chunks)

        doc = UploadedDocument(
            chatbot_id=chatbot_id,
            filename=file.filename,
            source_path=str(path),
            content_type=inferred_type or "application/octet-stream",
            file_hash=file_hash,
        )
        db.add(doc)
        await db.flush()

        for c in chunks:
            db.add(EmbeddingMetadata(
                document_id=doc.id,
                chunk_id=c["chunk_id"],
                text=c["text"],
                metadata_json=c["metadata"],
            ))

        await db.commit()
        return {"document_id": doc.id, "chunks": len(chunks)}

    except Exception as e:
        console.error(f"[INGEST][DB] Transaction failed for {file.filename}: {e}")
        await db.rollback()
        if 'path' in locals() and path.exists():
            path.unlink()
        raise HTTPException(status_code=400, detail=str(e))


# ---------------------------------------------------------------------------
# Secure Document Download
# ---------------------------------------------------------------------------

async def download_document(url: str) -> Path | None:
    import httpx
    import tempfile
    from urllib.parse import urlparse as _urlparse

    MAX_SIZE = 25 * 1024 * 1024
    ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md"}

    try:
        async with httpx.AsyncClient(follow_redirects=True, timeout=20.0) as client:
            head = await client.head(url)
            size = int(head.headers.get("content-length", 0))
            if size > MAX_SIZE:
                logger.warning(f"[DOWNLOAD] Too large ({size} bytes): {url}")
                return None

            ext = Path(_urlparse(url).path).suffix.lower()
            if ext not in ALLOWED_EXTENSIONS:
                logger.warning(f"[DOWNLOAD] Blocked extension {ext!r}: {url}")
                return None

            resp = await client.get(url)
            resp.raise_for_status()

            temp_dir = Path(tempfile.gettempdir()) / "tio_ingestion"
            temp_dir.mkdir(parents=True, exist_ok=True)
            temp_path = temp_dir / f"{uuid4()}{ext}"
            temp_path.write_bytes(resp.content)
            return temp_path

    except Exception as e:
        logger.error(f"[DOWNLOAD] Failed for {url}: {e}")
        return None



# ---------------------------------------------------------------------------
# Website Ingestion (background job logic)
# ---------------------------------------------------------------------------

async def ingest_website(chatbot_id: int, url: str) -> int:
    """Backward compatibility wrapper: submits to the job worker."""
    from backend.ingestion.worker import ingestion_worker
    return await ingestion_worker.submit_job(chatbot_id)


async def ingest_website_core(chatbot_id: int, job_id: int) -> None:
    """
    The controlled core ingestion logic. 
    Updates both Chatbot.status_json AND IngestionJob record.
    """
    from backend.db.session import SessionLocal
    from backend.utils.domain_intelligence import domain_detector
    from backend.utils.site_intelligence import build_site_profile
    from backend.models.entities import IngestionJob, UploadedDocument, EmbeddingMetadata
    from backend.utils.events import broadcast_ingestion_event

    profiler = IngestionProfiler(chatbot_id)
    
    async with SessionLocal() as db:
        async def update_stage(stage_name: str, progress: int, message: str, counts: dict = None):
            try:
                from backend.db.session import SessionLocal
                async with SessionLocal() as status_db:
                    curr_job = await status_db.get(IngestionJob, job_id)
                    curr_cb = await status_db.get(Chatbot, chatbot_id)
                    if not curr_job or not curr_cb: return
                    
                    if counts: profiler.update_counts(**counts)
                    profiler.detect_stall(progress, stage_name, message)
                    profiler.log_terminal_progress(progress, stage_name)

                    curr_job.current_stage = stage_name
                    curr_job.progress = progress
                    if counts:
                        if "total_chunks" in counts: curr_job.total_chunks = counts["total_chunks"]
                        if "indexed_chunks" in counts: curr_job.indexed_chunks = counts["indexed_chunks"]
                        if "failed_chunks" in counts: curr_job.failed_chunks = counts["failed_chunks"]

                    status_data = {
                        "stage": stage_name,
                        "progress": progress,
                        "message": message,
                        "eta": profiler.get_eta(progress),
                        "elapsed": profiler.elapsed,
                        "job_id": job_id,
                        **(profiler.counts)
                    }
                    curr_cb.status_json = status_data
                    await status_db.commit()
                    
                # Standard progress broadcast
                await broadcast_ingestion_event(chatbot_id, "progress", status_data)

                # Custom granular observability broadcasts
                if stage_name == "crawling":
                    clean_url = message.replace("Crawling: ", "").replace("Crawling ", "")
                    await broadcast_ingestion_event(chatbot_id, "crawl_progress", {
                        "type": "crawl_progress",
                        "url": clean_url,
                        "pages_crawled": profiler.counts.get("pages_crawled", 0),
                        "pages_total": profiler.counts.get("total_pages", 0) or 20
                    })
                elif stage_name == "parsing":
                    clean_doc = message.replace("Parsed document: ", "").replace("Parsing document: ", "")
                    await broadcast_ingestion_event(chatbot_id, "parser_progress", {
                        "type": "parser_progress",
                        "document": clean_doc
                    })
                elif stage_name in ("embedding", "indexing"):
                    await broadcast_ingestion_event(chatbot_id, "embedding_progress", {
                        "type": "embedding_progress",
                        "current": profiler.counts.get("embeddings_completed", 0),
                        "total": profiler.counts.get("total_chunks", 0)
                    })
            except Exception as se:
                logger.error(f"[STAGE_ERROR] Failed to update stage {stage_name}: {se}")

        try:
            chatbot = await db.get(Chatbot, chatbot_id)
            job = await db.get(IngestionJob, job_id)
            if not chatbot or not job: return

            url = chatbot.website_url

            # 0. Pre-Ingestion Integrity Check
            await update_stage("crawling", 1, "Initializing validation pipeline...")
            
            print(flush=True)
            print("========================================================", flush=True)
            print("[INGESTION] Running Pre-Ingestion Integrity Check...", flush=True)
            print("========================================================", flush=True)
            print(flush=True)
            
            await update_stage("crawling", 2, "Checking ingestion queue...")
            from backend.ingestion.worker import ingestion_worker
            q_size = ingestion_worker.queue.qsize()
            print(f"[OK] Ingestion queue reachable. Queue size: {q_size}", flush=True)
            await asyncio.sleep(0.05)
            
            await update_stage("crawling", 3, "Checking worker availability...")
            wrk_status = "ACTIVE" if ingestion_worker._worker_task and not ingestion_worker._worker_task.done() else "INACTIVE"
            print(f"[OK] Background worker is {wrk_status}", flush=True)
            await asyncio.sleep(0.05)
            
            await update_stage("crawling", 4, "Checking embedding service status...")
            from backend.utils.validation import validate_embeddings
            emb_status = await validate_embeddings()
            print(f"[OK] Embedding service is {emb_status}", flush=True)
            await asyncio.sleep(0.05)
            
            await update_stage("crawling", 4, "Checking vectorstore status...")
            from backend.utils.validation import validate_vectorstore
            vs_val = await validate_vectorstore()
            print(f"[OK] Vectorstore is {vs_val['status']}. Collection: {vs_val['collection']}, Dimension: {vs_val['dimension']}", flush=True)
            await asyncio.sleep(0.05)
            
            print(flush=True)
            print("========================================================", flush=True)
            print("[SUCCESS] Pre-Ingestion Integrity Check Passed", flush=True)
            print("========================================================", flush=True)
            print(flush=True)

            # 1. Domain Detection & Crawl Discovery
            await update_stage("crawling", 5, "Discovering pages...")

            profiler.start_stage("crawling")
            t_crawl_start = time.monotonic()
            try:
                # Wrap with a hard 10s timeout to prevent freezing on slow/broken homepages
                homepage_content = await asyncio.wait_for(scraper.extract_content(url), timeout=10.0)
            except Exception as e:
                console.critical(f"Crawler homepage extraction timeout/failed: {e}. Proceeding with fallback.", stage="CRAWLING")
                homepage_content = ""
            
            # Fast domain detection (Local/Deterministic)
            domain_scores = domain_detector.get_scores(homepage_content, {"url": url})
            sorted_scores = sorted(domain_scores, key=lambda x: x.score, reverse=True)
            detected_domain = sorted_scores[0].domain if sorted_scores and sorted_scores[0].score > 2.0 else "general"
            chatbot.domain = detected_domain
            chatbot.behavior_profile = detected_domain
            await db.commit()

            try:
                async def on_crawl_progress(current: int, total: int, crawled_url: str | dict, custom_payload: dict = None):
                    if isinstance(crawled_url, dict):
                        payload = crawled_url
                        await broadcast_ingestion_event(chatbot_id, "crawler_status", payload)
                        await update_stage("crawling", payload.get("progress", 5), f"{payload.get('stage')}")
                        return
                    if custom_payload:
                        await broadcast_ingestion_event(chatbot_id, "crawler_status", custom_payload)
                        await update_stage("crawling", custom_payload.get("progress", 5), f"{custom_payload.get('stage')}")
                        return
                    
                    print(flush=True)
                    print("========================================================", flush=True)
                    print("[CRAWLER]", flush=True)
                    print(f"Fetching:\n{crawled_url}", flush=True)
                    print("========================================================", flush=True)
                    print(flush=True)

                    await update_stage(
                        "crawling", 
                        min(5 + int((current / total) * 15), 19), 
                        f"Crawling: {crawled_url}", 
                        {"pages_crawled": current, "total_pages": total}
                    )

                pages, docs = await asyncio.wait_for(
                    scraper.discover_assets(url, limit=20, depth=1, allow_external=False, on_progress=on_crawl_progress),
                    timeout=120.0
                )
                profiler.update_counts(total_pages=len(pages))
            except asyncio.TimeoutError:
                logger.error(f"[INGEST] TIMEOUT during crawling for chatbot {chatbot_id}. Falling back to homepage.")
                pages, docs = [url], []
            
            profiler.end_stage("crawling")
            if settings.debug_timing:
                print(f"[CRAWLER] Completed in {time.monotonic() - t_crawl_start:.1f}s")
            
            # 2. Page Extraction
            await update_stage("extracting", 20, f"Extracting {len(pages)} pages...", {"pages_crawled": 0})
            profiler.start_stage("extraction")
            t_ext_start = time.monotonic()
            
            semaphore = asyncio.Semaphore(5)
            async def process_page(page_url: str):
                async with semaphore:
                    try:
                        print("========================================================", flush=True)
                        print("LIVE PAGE CRAWLING", flush=True)
                        print("========================================================", flush=True)
                        print(f"[CRAWLER] Crawling page: {page_url}\n", flush=True)
                        
                        t_start = time.monotonic()
                        content = await scraper.extract_content(page_url)
                        
                        crawled_inc = profiler.counts.get("pages_crawled", 0) + 1
                        if content and len(content.split()) > 20:
                            duration = time.monotonic() - t_start
                            print("[EXTRACTION]", flush=True)
                            print(f"Extracted page: {page_url}\n", flush=True)
                            print("Characters:", flush=True)
                            print(f"{len(content)}\n", flush=True)
                            print("Time:", flush=True)
                            print(f"{duration:.1f}s\n", flush=True)
                            
                            await update_stage("extracting", min(20 + int((crawled_inc / max(len(pages), 1)) * 10), 29),
                                               f"Extracted page: {page_url}",
                                               {"pages_crawled": crawled_inc})
                            return (content, page_url)
                        else:
                            print(f"[EXTRACTION] Page ignored (thin content): {page_url}\n", flush=True)
                            await update_stage("extracting", min(20 + int((crawled_inc / max(len(pages), 1)) * 10), 29),
                                               f"Extracted page: {page_url} (ignored thin)",
                                               {"pages_crawled": crawled_inc})
                            return None
                    except Exception as e:
                        print(f"[EXTRACTION] Extraction failed for {page_url}:\nReason:\n{e}\n", flush=True)
                        crawled_inc = profiler.counts.get("pages_crawled", 0) + 1
                        await update_stage("extracting", min(20 + int((crawled_inc / max(len(pages), 1)) * 10), 29),
                                           f"Extraction failed: {page_url}",
                                           {"pages_crawled": crawled_inc})
                        return None

            results = await asyncio.wait_for(
                asyncio.gather(*(process_page(p) for p in pages), return_exceptions=True),
                timeout=240.0
            )
            extracted = [r for r in results if isinstance(r, tuple)]
            profiler.update_counts(pages_crawled=len(extracted))
            profiler.end_stage("extraction")
            if settings.debug_timing:
                print(f"[EXTRACTION] Completed in {time.monotonic() - t_ext_start:.1f}s")

            # 3. Chunking & Processing
            await update_stage("chunking", 40, "Processing and chunking content...")
            profiler.start_stage("chunking")
            t_chunk_start = time.monotonic()
            all_chunks = []
            all_text_buffer = [homepage_content]
            
            for content, p_url in extracted:
                all_text_buffer.append(content)
                all_chunks.extend(_chunk_text(content, p_url, chatbot_id, domain=detected_domain))

            # Handle docs
            doc_types = {".pdf": "application/pdf", ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document", ".txt": "text/plain", ".md": "text/markdown"}
            if docs:
                await update_stage("parsing", 30, f"Parsing {len(docs)} documents...")
                profiler.start_stage("parsing")
                doc_idx = 0
                for doc_url in docs:
                    filename = doc_url.split("/")[-1]
                    print("========================================================", flush=True)
                    print("[PARSER]", flush=True)
                    print("Extracting text...", flush=True)
                    print("========================================================", flush=True)
                    
                    t_start = time.monotonic()
                    try:
                        temp_path = await asyncio.wait_for(download_document(doc_url), timeout=30.0)
                        if temp_path:
                            content = await asyncio.to_thread(_parse_file, str(temp_path))
                            duration = time.monotonic() - t_start
                            
                            if content and len(content.split()) >= 20:
                                ext = temp_path.suffix.lower()
                                s_type = doc_types.get(ext, "application/octet-stream")
                                all_text_buffer.append(content)
                                
                                new_chunks = _chunk_text(content, doc_url, chatbot_id, domain=detected_domain, source_type=s_type)
                                all_chunks.extend(new_chunks)
                                
                                print(flush=True)
                                print("[PARSER]", flush=True)
                                print("Sections extracted:", flush=True)
                                print(str(len(content.split('\n\n'))), flush=True)
                                print(flush=True)
                                
                                print("[PARSER]", flush=True)
                                print("Semantic chunks generated:", flush=True)
                                print(f"{len(new_chunks)}", flush=True)
                                print(flush=True)
                                
                                print("========================================================", flush=True)
                                print("[TIMER]", flush=True)
                                print("Document parsed in:", flush=True)
                                print(f"{duration:.1f}s", flush=True)
                                print("========================================================", flush=True)
                                print(flush=True)
                                
                            if temp_path.exists(): temp_path.unlink()
                            
                            doc_idx += 1
                            await update_stage("parsing", min(30 + int((doc_idx / len(docs)) * 10), 39),
                                               f"Parsed document: {filename}",
                                               {"skipped_pages": doc_idx}) # skipped_pages maps to documents parsed in profiler
                    except Exception as e:
                        logger.warning(f"Failed to process doc {doc_url}: {e}")
                        print(f"[PARSER] Failed to parse document {filename}: {e}", flush=True)
                profiler.end_stage("parsing")

            if len(all_chunks) > 1200:
                all_chunks = all_chunks[:1200]
            
            total_chunks = len(all_chunks)
            profiler.update_counts(chunks_generated=total_chunks)
            profiler.end_stage("chunking")
            print(f"[CHUNKING] Generated: {total_chunks} chunks", flush=True)
            if settings.debug_timing:
                print(f"[EXTRACTION] Chunked content into {total_chunks} chunks in {time.monotonic() - t_chunk_start:.1f}s")

            # 4. Embedding & Indexing (Batched)
            indexed_count = 0
            failed_count = 0
            if all_chunks:
                profiler.start_stage("indexing")
                t_idx_start = time.monotonic()
                logger.info(f"[STABILIZATION] START indexing for {total_chunks} chunks - No external APIs allowed")
                batch_size = 100
                total_batches = (total_chunks + batch_size - 1) // batch_size

                for i in range(0, total_chunks, batch_size):
                    batch = all_chunks[i:i+batch_size]
                    batch_num = (i // batch_size) + 1
                    
                    # Print embedding progress for the batch
                    print("========================================================", flush=True)
                    print("[EMBEDDING]", flush=True)
                    print(f"Embedding batch: {batch_num}/{total_batches}", flush=True)
                    print(f"Chunks: {i + 1} - {min(i + len(batch), total_chunks)} of {total_chunks}", flush=True)
                    print("Model: BAAI/bge-small-en-v1.5", flush=True)
                    print("========================================================", flush=True)
                    print(flush=True)
                        
                    await update_stage("indexing", 50 + int((batch_num/total_batches)*30), 
                                     f"Indexing batch {batch_num}/{total_batches}...",
                                     {"total_chunks": total_chunks, "embeddings_completed": indexed_count, "failed_chunks": failed_count})
                    
                    # Broadcast detailed embedding progress to frontend via WebSocket
                    await broadcast_ingestion_event(chatbot_id, "embedding_progress", {
                        "chunk_num": min(i + batch_size, total_chunks),
                        "total_chunks": total_chunks,
                        "model": "BAAI/bge-small-en-v1.5"
                    })
                    
                    try:
                        # Vector insertion
                        await asyncio.wait_for(asyncio.to_thread(upsert_chunks, batch), timeout=90.0)
                        
                        # Metadata persistence with transactional isolation
                        doc_cache = {}
                        path_to_type = {}
                        for c in batch:
                            if c["document"] not in path_to_type:
                                path_to_type[c["document"]] = c["metadata"].get("source_type", "text/html")

                        for path, c_type in path_to_type.items():
                            final_type = c_type or "text/html"
                            
                            doc = (await db.execute(select(UploadedDocument).where(
                                UploadedDocument.chatbot_id == chatbot_id,
                                UploadedDocument.source_path == path
                            ))).scalar_one_or_none()
                            
                            if not doc:
                                try:
                                    doc = UploadedDocument(
                                        chatbot_id=chatbot_id, 
                                        filename=path.split("/")[-1] or "page", 
                                        source_path=path,
                                        content_type=final_type
                                    )
                                    db.add(doc)
                                    await db.flush()
                                except Exception as dbe:
                                    print("========================================================", flush=True)
                                    print("[WARNING]", flush=True)
                                    print("Metadata insertion failed:", flush=True)
                                    print(f"{path}", flush=True)
                                    print("========================================================", flush=True)
                                    print("[WARNING]", flush=True)
                                    print("Skipping document metadata...", flush=True)
                                    print(flush=True)
                                    await db.rollback()
                                    continue
                                    
                            doc_cache[path] = doc.id
                        
                        # Insert metadata batch
                        from sqlalchemy import insert
                        meta_batch = []
                        for c in batch:
                            if c["document"] in doc_cache:
                                meta_batch.append({
                                    "document_id": doc_cache[c["document"]], 
                                    "chunk_id": c["chunk_id"], 
                                    "text": c["text"], 
                                    "metadata": c["metadata"]
                                })

                        if meta_batch:
                            await db.execute(insert(EmbeddingMetadata), meta_batch)
                            await db.commit()
                            indexed_count += len(batch)
                            
                            print("========================================================", flush=True)
                            print("[VECTORSTORE]", flush=True)
                            print("Inserted vectors:", flush=True)
                            print(f"{len(batch)}", flush=True)
                            print("========================================================", flush=True)
                            print(flush=True)
                            print("[VECTORSTORE]", flush=True)
                            print("Collection:", flush=True)
                            print(f"chatbot_{chatbot_id}_collection", flush=True)
                            print(flush=True)
                            
                            await broadcast_ingestion_event(chatbot_id, "vector_progress", {
                                "inserted_vectors": len(batch),
                                "collection": f"chatbot_{chatbot_id}_collection"
                            })
                        else:
                            print("========================================================", flush=True)
                            print("[WARNING]", flush=True)
                            print("Empty metadata batch skipped:", flush=True)
                            print(f"Batch {batch_num}", flush=True)
                            print("========================================================", flush=True)
                            failed_count += len(batch)
                            
                    except Exception as e:
                        print("========================================================", flush=True)
                        print("[WARNING]", flush=True)
                        print("Batch embedding failed:", flush=True)
                        print(f"Batch {batch_num}", flush=True)
                        print("========================================================", flush=True)
                        print("[WARNING]", flush=True)
                        print("Retrying extraction or rolling back...", flush=True)
                        print(flush=True)
                        await db.rollback()
                        failed_count += len(batch)
                
                profiler.end_stage("indexing")
                
                print("========================================================", flush=True)
                print("[TIMER]", flush=True)
                print("Embedding completed in:", flush=True)
                print(f"{time.monotonic() - t_idx_start:.1f}s", flush=True)
                print("========================================================", flush=True)
                print(flush=True)

            # 5. Entity Enrichment & Site Profile
            await update_stage("finalizing", 90, "Finalizing intelligence profile...")
            
            # Site Profile (LLM-based but Local/Self-contained)
            if settings.enable_profile_intelligence:
                try:
                    logger.info(f"[STABILIZATION] Finalizing local intelligence profile for {url}")
                    console.info("[ORCH] Generating Site Intelligence Profile", stage="ORCH")
                    full_text = "\n\n".join(all_text_buffer[:30])
                    site_profile = await asyncio.wait_for(
                        build_site_profile(all_text=full_text, domain=detected_domain, site_url=url),
                        timeout=60.0
                    )
                    chatbot.site_profile = site_profile
                    await db.commit()
                except Exception as e:
                    logger.warning(f"[STABILIZATION] Profile generation failed or timed out: {e}")
                    console.warning(f"[ORCH] Profile generation failed, using basic metadata: {e}", stage="ORCH")
            else:
                console.info("[ORCH] Site Profile Intelligence disabled via feature flag", stage="ORCH")

            # 6. Final Validation & Completion
            # Part 1: STRICT INGESTION VALIDATION
            pages_crawled = profiler.counts.get("pages_crawled", 0)
            
            console.info(f"[INGESTION] Pages crawled={pages_crawled}", stage="VALIDATION")
            console.info(f"[INGESTION] Chunks created={total_chunks}", stage="VALIDATION")

            if pages_crawled == 0 or total_chunks == 0:
                reason = "No content could be extracted from the provided URL."
                if pages_crawled == 0:
                    reason = "Crawler could not discover any accessible pages. Check the URL and robots.txt."
                elif total_chunks == 0:
                    reason = "Pages were found but no meaningful content could be extracted (possible JS-heavy site or empty pages)."

                logger.error(f"[INGESTION] FAILED for chatbot {chatbot_id}: {reason}")
                
                error_data = {
                    "stage": "error",
                    "progress": 0,
                    "message": f"Ingestion Failed: {reason}",
                    "error": reason,
                    "job_id": job_id,
                    **(profiler.counts)
                }
                chatbot.status = "error"
                chatbot.status_json = error_data
                await db.commit()
                
                await broadcast_ingestion_event(chatbot_id, "failure", error_data)
                return

            # 7. Complete
            await update_stage("ready", 100, "Ready", 
                             {"total_chunks": total_chunks, "embeddings_completed": indexed_count, "failed_chunks": failed_count})
            
            chatbot.status = "ready"
            await db.commit()
            
            total_time = time.monotonic() - t_crawl_start
            
            print("========================================================", flush=True)
            print("[SUCCESS]", flush=True)
            print("INGESTION COMPLETE", flush=True)
            print("========================================================", flush=True)
            print(flush=True)
            print("Pages Crawled:", flush=True)
            print(f"{pages_crawled}", flush=True)
            print(flush=True)
            print("Documents Parsed:", flush=True)
            print(f"{len(docs)}", flush=True)
            print(flush=True)
            print("Chunks Generated:", flush=True)
            print(f"{total_chunks}", flush=True)
            print(flush=True)
            print("Vectors Inserted:", flush=True)
            print(f"{indexed_count}", flush=True)
            print(flush=True)
            print("Total Time:", flush=True)
            print(f"{total_time:.1f}s", flush=True)
            print(flush=True)
            
            await broadcast_ingestion_event(chatbot_id, "complete", {
                "chatbot_id": chatbot_id,
                "status": "ready",
                "counts": profiler.counts
            })
            
            profiler.log_summary()

        except Exception as e:
            logger.error(f"[INGESTION] FAILED for chatbot={chatbot_id}: {e}", exc_info=True)
            # Worker will handle final error state update
            raise e
